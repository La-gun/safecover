# -*- coding: utf-8 -*-
"""Build Word partner pitch (run from repo root: python docs/partner/build_partner_pitch_docx.py)."""
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
OUT = HERE / "SafeCover-Insurance-Partner-Pitch.docx"

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Install: pip install python-docx", file=sys.stderr)
    raise

from pitch_assets import render_all_diagram_pngs, render_chart_pngs


def _p(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    p.paragraph_format.space_after = Pt(6)
    return p


def _bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def _add_picture(doc, buf, width_in=6.2):
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width_in))
    doc.add_paragraph()


def build():
    doc = Document()
    core = doc.core_properties
    core.title = "SafeCover — Insurance Partner Pitch"
    core.subject = "Business case, architecture, partnership path"
    core.keywords = "SafeCover, embedded insurance, microinsurance, POS"

    t = doc.add_heading("SafeCover", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = doc.add_paragraph(
        "Embedded microinsurance for e-commerce and retail POS: business case, architecture, "
        "and partnership path for a licensed insurer."
    )
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in st.runs:
        r.font.size = Pt(12)
    doc.add_paragraph(
        "Confidential — partner discussion. Prepared May 2026. Builds on the internal white paper draft "
        "(Embedded MicroInsurance Platform Business Model) and the SafeCover reference implementation."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Executive summary", level=1)
    _p(
        doc,
        "SafeCover is an API-first embedded microinsurance orchestration layer: quotes, bind, "
        "payment-aligned confirmation, webhooks, and a unified retail POS flow—so merchants can offer "
        "contextual protection without sending shoppers to a separate insurance journey.",
    )
    _p(
        doc,
        "Ask to the carrier: pilot one filed product (or rider) through SafeCover’s integration surface, "
        "connect your PAS and claims rails, and co-sell through one anchor distributor (e-commerce or POS)—"
        "proving attach rate, loss behaviour, and operational fit in 90–120 days.",
        bold=True,
    )

    doc.add_heading("Market opportunity", level=1)
    _p(
        doc,
        "Embedded insurance shifts purchase from “find a policy later” to “protect this basket, trip, or ticket now.” "
        "Industry commentary points to a large gross-premium pool this decade; the strategic direction is consistent: "
        "ecosystems win distribution, and carriers win when partners meet them with APIs, speed, and clear compliance hooks.",
    )
    _p(
        doc,
        "Microinsurance fits short, well-defined exposures. SafeCover uses scenario-based offers and multi-provider "
        "comparison in demo workshops, with a production path to a single approved carrier per market.",
    )
    for buf in render_chart_pngs():
        _add_picture(doc, buf, 6.3)
    _p(
        doc,
        "Illustrative charts for discussion only—not firm forecasts or scored research.",
        bold=False,
    )

    doc.add_heading("Stakeholder value", level=1)
    for buf in render_all_diagram_pngs():
        _add_picture(doc, buf, 6.3)

    doc.add_heading("What SafeCover is today (reference stack)", level=1)
    doc.add_heading("Distribution surface", level=2)
    _bullet(doc, "Node.js Express API: quote, bind, confirm, webhook, claims stubs")
    _bullet(doc, "POST /api/pos/enhanced for retail: quote, bind, confirm, or sale")
    _bullet(doc, "Universal widget + adapters: Shopify, WooCommerce, BigCommerce, Magento, generic")
    _bullet(doc, "WooCommerce PHP plugin path for server-side checkout")
    doc.add_heading("Insurer-grade hooks (prototype)", level=2)
    _bullet(doc, "Strict mode: API keys, quote signing, registered quotes at bind, CORS allowlist")
    _bullet(doc, "Webhook HMAC, bind idempotency, regulatory_snapshot on policies")
    _bullet(doc, "Environment placeholders: carrier entity, license ref, filing code, wording version")
    _bullet(doc, "Solidity contract in repo; on-chain issuance from API is a roadmap item")
    _p(
        doc,
        "See README.md, docs/API.md, docs/PRODUCTION.md, and docs/WHITEPAPER-CODEBASE-MAP.md in the repository.",
    )

    doc.add_heading("Business case for the insurer", level=1)
    _bullet(doc, "Distribution leverage: digital-native volumes through checkout and POS vs standalone D2C acquisition.")
    _bullet(doc, "Product modularisation: scenario-tariff modules map to filed ingredients you approve.")
    _bullet(doc, "Operational efficiency: server-to-server APIs, idempotent binds, webhook contracts.")
    _bullet(doc, "Innovation narrative: embedded distribution with optional ledger path where appropriate.")

    doc.add_heading("Proposed pilot (90–120 days)", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Phase", "Outcomes"),
        ("1 · Design", "Chosen scenario(s), tariff mapping, disclosures, sandbox PAS events."),
        ("2 · Integrate", "Strict mode in staging; webhook to PAS; confirm on payment capture."),
        ("3 · Live limited", "One merchant or franchise POS cohort; monitoring, loss run, UX metrics."),
        ("4 · Scale decision", "Attach rate, conversion, claims cycle, partner revenue—go / refine / pause."),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b

    doc.add_heading("Transparency (underwriting trust)", level=1)
    _p(
        doc,
        "The reference implementation does not replace filings, capital, licensed claims payment, or full PAS depth. "
        "Items called out in WHITEPAPER-CODEBASE-MAP.md as not yet wired include on-chain issuance from the API, "
        "parametric oracle payouts, and full carrier billing integration—these are partnership build scope, with "
        "SafeCover supplying the distribution orchestration skeleton.",
    )

    doc.add_heading("References (from your draft)", level=1)
    _p(
        doc,
        "BCG embedded insurance tech stack; industry surveys on checkout protection; case patterns (airline attach, "
        "marketplace warranty uplift); Etherisc / parametric microinsurance references—as cited in your original Word document. "
        "Use compliance-approved third-party materials where required.",
    )

    doc.add_heading("Regeneration", level=1)
    _p(
        doc,
        "HTML/PDF: python docs/partner/build_partner_pitch_html.py. Word: python docs/partner/build_partner_pitch_docx.py. "
        "Requires matplotlib and python-docx.",
    )

    doc.save(OUT)
    print("Wrote", OUT, "bytes", OUT.stat().st_size)


if __name__ == "__main__":
    build()
